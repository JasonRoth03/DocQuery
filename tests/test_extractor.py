import pytest
from io import BytesIO
from app.services.extractor import extract_text


def test_plain_text():
    data = b"Hello, world!\nThis is a test."
    assert extract_text(data, "text/plain", "test.txt") == "Hello, world!\nThis is a test."


def test_utf8_characters():
    text = "Héllo wörld — “smart quotes”"
    result = extract_text(text.encode("utf-8"), "text/plain", "test.txt")
    assert result == text


def test_markdown_by_extension():
    data = b"# Header\n\nSome content here."
    result = extract_text(data, "application/octet-stream", "README.md")
    assert "Header" in result
    assert "content" in result


def test_txt_by_extension_overrides_generic_mime():
    data = b"plain text content"
    result = extract_text(data, "application/octet-stream", "notes.txt")
    assert result == "plain text content"


def test_docx():
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.add_paragraph("")  # blank paragraph — should be skipped

    buf = BytesIO()
    doc.save(buf)

    result = extract_text(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "test.docx",
    )
    assert "First paragraph" in result
    assert "Second paragraph" in result
    # blank paragraph should not appear as a lone newline cluster
    assert result.strip() == "First paragraph\n\nSecond paragraph"


def test_unsupported_mime_raises():
    with pytest.raises(ValueError, match="Unsupported"):
        extract_text(b"\x89PNG\r\n", "image/png", "photo.png")


def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="Unsupported"):
        extract_text(b"data", "application/octet-stream", "file.xyz")
